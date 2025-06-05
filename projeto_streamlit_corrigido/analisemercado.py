#!/usr/bin/env python
# coding: utf-8

import os
import re
import uuid
import logging
import markdown
import matplotlib
matplotlib.use('Agg') # Use Agg backend for non-interactive environments like Render
import matplotlib.pyplot as plt
from datetime import datetime
# from dotenv import load_dotenv # Removed: Load variables directly from environment in production (e.g., Render)
from crewai import Agent, Task, Crew
from crewai_tools.serper_dev_tool import SerperDevTool
from crewai_tools.scrape_website_tool import ScrapeWebsiteTool
from crewai_tools.file_writer_tool import FileWriterTool
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE

# Configuração de logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Carregar variáveis de ambiente (REMOVIDO PARA PRODUÇÃO)
# load_dotenv() # Ensure API keys (OPENAI_API_KEY, SERPER_API_KEY, GROQ_API_KEY) are set as environment variables in Render

# Ferramentas (CrewAI tools typically load keys from environment variables if not provided explicitly)
# Ensure SERPER_API_KEY is set in the environment
serper = SerperDevTool()
scraper = ScrapeWebsiteTool()
file_writer = FileWriterTool()

# Configuração do relatório
CONFIG_RELATORIO = {
    "paginas_minimas": 10,
    "contagem_palavras": 10000,
    "secoes": [
        "Descrição da Categoria", "Panorama de Mercado", "Análise da Cadeia de Suprimentos",
        "Indicadores Econômicos", "Benchmarking", "Análise SWOT", "5 Forças de Porter", "CBD (Cost Breakdown)",
        "SCA (Should Cost Analysis)", "LPP (Line Performance Pricing)", "Perguntas para Clientes e Fornecedores",
        "SLAs e Multas", "Critérios de Seleção de Fornecedores", "Estrutura de RFP", "Análise de Riscos",
        "Alavancas de Negociação e BATNA", "Momento Ideal para Negociação", "Tendências ESG", "Conclusão"
    ]
}

# Função para criar gráfico aranha
def criar_grafico_aranha(dados, rotulos, titulo, arquivo):
    try:
        if not dados or not rotulos or len(dados) != len(rotulos):
            logging.warning(f"Dados ou rótulos inválidos para o gráfico '{titulo}'. Gráfico não gerado.")
            return False

        angulos = [n / float(len(rotulos)) * 2 * 3.14159 for n in range(len(rotulos))]
        angulos += angulos[:1]
        valores = list(dados.values()) + list(dados.values())[:1]

        fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
        ax.fill(angulos, valores, color='blue', alpha=0.25)
        ax.plot(angulos, valores, color='blue', linewidth=2) # Add outline
        ax.set_xticks(angulos[:-1])
        ax.set_xticklabels(rotulos)
        ax.set_yticks([]) # Hide radial axis labels/ticks if not needed
        ax.set_title(titulo, size=12, y=1.1)
        plt.savefig(arquivo, format='png', bbox_inches='tight')
        plt.close(fig) # Close the figure to free memory
        logging.info(f"Gráfico salvo como {arquivo}")
        return True
    except Exception as e:
        logging.error(f"Erro ao criar gráfico aranha '{titulo}': {e}", exc_info=True)
        return False

# Função auxiliar para adicionar tabela Markdown no Word
def adicionar_tabela_md_no_word(linhas_md, doc):
    try:
        # Remove a linha de separador (|---|---|)
        linhas = [linha for linha in linhas_md if not re.match(r'^\|[-\s|]+\|$', linha)]
        if not linhas or len(linhas) < 2: # Precisa de cabeçalho e pelo menos uma linha de dados
            logging.warning("Dados de tabela Markdown insuficientes ou inválidos.")
            return

        # Extrai cabeçalho
        header = [cell.strip() for cell in linhas[0].strip('|').split('|')]
        if not header:
            logging.warning("Cabeçalho da tabela Markdown vazio.")
            return

        tabela = doc.add_table(rows=1, cols=len(header))
        tabela.style = 'Table Grid' # Apply basic grid style
        tabela.autofit = True
        hdr_cells = tabela.rows[0].cells
        for i, texto in enumerate(header):
            hdr_cells[i].text = texto
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True # Make header bold

        # Adiciona linhas de dados
        for linha_dados in linhas[1:]:
            valores = [cell.strip() for cell in linha_dados.strip('|').split('|')]
            # Ensure number of values matches number of columns
            if len(valores) != len(header):
                logging.warning(f"Número de colunas inconsistente na linha: {linha_dados}. Esperado: {len(header)}, Obtido: {len(valores)}")
                # Pad or truncate to match header length
                valores.extend([''] * (len(header) - len(valores))) # Pad if too short
                valores = valores[:len(header)] # Truncate if too long

            row_cells = tabela.add_row().cells
            for i, texto in enumerate(valores):
                 row_cells[i].text = texto

    except Exception as e:
        logging.error(f"Erro ao adicionar tabela Markdown ao Word: {e}", exc_info=True)

# Função para exportar conteúdo Markdown para documento Word
def exportar_para_word(conteudo_md, setor, arquivo_saida):
    try:
        doc = Document()
        # Basic styles (can be customized further)
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        doc.add_heading(f'Relatório de Análise de Mercado: {setor}', level=0)
        doc.add_paragraph('Preparado por Vorätte Consultoria', style='Subtitle') # Use built-in subtitle style
        doc.add_paragraph() # Add a blank paragraph for spacing

        # Define custom styles if they don't exist
        if 'TituloNegrito1' not in doc.styles:
            style = doc.styles.add_style('TituloNegrito1', WD_STYLE_TYPE.PARAGRAPH)
            style.font.bold = True
            style.font.size = Pt(14)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        if 'TituloNegrito2' not in doc.styles:
            style = doc.styles.add_style('TituloNegrito2', WD_STYLE_TYPE.PARAGRAPH)
            style.font.bold = True
            style.font.size = Pt(12)
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(5)

        linhas = conteudo_md.split('\n')
        i = 0
        in_code_block = False
        while i < len(linhas):
            linha = linhas[i] # Keep leading/trailing whitespace for now

            # Handle code blocks
            if linha.strip().startswith('```'):
                in_code_block = not in_code_block
                i += 1
                continue
            if in_code_block:
                # Optionally add code block content with a specific style
                # para = doc.add_paragraph(linha, style='CodeStyle') # Requires defining 'CodeStyle'
                i += 1
                continue

            linha_strip = linha.strip()

            # Skip empty lines and separators
            if not linha_strip or linha_strip == '---':
                i += 1
                continue

            # Skip markdown images
            if re.match(r'^!\[.*\]\(.*\)', linha_strip):
                logging.info(f"Ignorando imagem Markdown: {linha_strip}")
                i += 1
                continue

            # Handle tables
            if linha_strip.startswith('|') and '|' in linha_strip:
                tabela_linhas = []
                start_table_index = i
                while i < len(linhas) and linhas[i].strip().startswith('|'):
                    tabela_linhas.append(linhas[i].strip())
                    i += 1
                logging.info(f"Encontrada tabela Markdown nas linhas {start_table_index}-{i-1}")
                adicionar_tabela_md_no_word(tabela_linhas, doc)
                doc.add_paragraph() # Add space after table
                continue # Skip incrementing i again

            # Handle headings
            if linha_strip.startswith('# '):
                doc.add_paragraph(linha_strip[2:].strip(), style='TituloNegrito1')
            elif linha_strip.startswith('## '):
                doc.add_paragraph(linha_strip[3:].strip(), style='TituloNegrito2')
            elif linha_strip.startswith('### '):
                 # Use TituloNegrito2 or create a TituloNegrito3 style
                doc.add_paragraph(linha_strip[4:].strip(), style='TituloNegrito2')
            # Handle list items (basic conversion)
            elif re.match(r'^\s*[-\*\+]\s+', linha_strip):
                texto_item = re.sub(r'^\s*[-\*\+]\s+', '', linha_strip)
                # Basic formatting removal (bold/italic)
                texto_item = re.sub(r'\*\*?(.*?)\*\*?', r'\1', texto_item)
                doc.add_paragraph(texto_item, style='List Bullet') # Use built-in list style
            # Handle regular paragraphs
            else:
                # Basic formatting removal (bold/italic, code ticks, links)
                linha_limpa = re.sub(r'\*\*?(.*?)\*\*?', r'\1', linha_strip)
                linha_limpa = re.sub(r'`{1,3}(.*?)`{1,3}', r'\1', linha_limpa)
                linha_limpa = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', linha_limpa) # Keep link text, remove URL part
                doc.add_paragraph(linha_limpa)

            i += 1

        # Add generated images (ensure paths are correct)
        # Using relative paths, assuming they are generated in the same directory
        swot_image_path = 'swot_aranha.png'
        porter_image_path = 'porter_aranha.png'

        if os.path.exists(swot_image_path):
            try:
                doc.add_picture(swot_image_path, width=Inches(4))
                logging.info(f"Imagem {swot_image_path} adicionada ao documento.")
            except Exception as img_e:
                 logging.error(f"Erro ao adicionar imagem {swot_image_path}: {img_e}")
        else:
            logging.warning(f"Arquivo de imagem {swot_image_path} não encontrado.")

        if os.path.exists(porter_image_path):
            try:
                doc.add_picture(porter_image_path, width=Inches(4))
                logging.info(f"Imagem {porter_image_path} adicionada ao documento.")
            except Exception as img_e:
                 logging.error(f"Erro ao adicionar imagem {porter_image_path}: {img_e}")
        else:
            logging.warning(f"Arquivo de imagem {porter_image_path} não encontrado.")

        doc.save(arquivo_saida)
        logging.info(f"Documento Word salvo como: {arquivo_saida}")

    except Exception as e:
        logging.error(f"Erro fatal ao exportar para Word: {e}", exc_info=True)
        raise # Re-raise the exception to be caught by the caller

# Agentes (Ensure API keys are set in environment variables)
# OPENAI_API_KEY, GROQ_API_KEY, SERPER_API_KEY
pesquisador = Agent(
    role="Pesquisador de Mercado",
    goal="Coletar dados completos sobre {setor} em {regiao} para {empresa}, cobrindo oferta/demanda, competidores, custos, forças, fortalezas, indicadores econômicos, ESG e riscos.",
    backstory="Analista experiente em sourcing estratégico, especializado em dados de mercado.",
    tools=[serper, scraper],
    verbose=True,
    allow_delegation=False # Example: Prevent delegation if not needed
)

analista = Agent(
    role="Analista de Tendências",
    goal="Analisar dados de {setor} para gerar insights, incluindo SWOT, 5 Forças de Porter, TCO, LPP, CBD, SCA e tendências ESG.",
    backstory="Especialista em estratégias de compras, negociação, focado em insights acionáveis.",
    tools=[serper],
    verbose=True,
    allow_delegation=False
)

redator = Agent(
    role="Redator de Relatórios",
    goal="Criar um relatório detalhado em Markdown sobre {setor}, com visualizações, exportado para Word.",
    backstory="Redator profissional que transforma dados complexos em relatórios estratégicos.",
    verbose=True,
    allow_delegation=False
)

# Tarefas
coleta_dados = Task(
    description=(
        "Coletar dados sobre {setor} em {regiao} para {empresa}, incluindo:\n"
        "1. Panorama de oferta e demanda (global e regional).\n"
        "2. Top 20 competidores por faturamento (2024 ou mais recente).\n"
        "3. Modelos de contratação (ex.: pay-as-you-go, reservado).\n"
        "4. Estrutura de custos e tendências de preços (últimos 24 meses, vs. USD, IGPM, IPCA).\n"
        "5. CBD, SCA, TCO, LPP, SLAs e multas.\n"
        "6. Detalhes da cadeia de suprimentos (produtores, logística, dependência de importação).\n"
        "7. Práticas e tendências ESG.\n"
        "8. Riscos (geopolíticos, regulatórios, etc.) e oportunidades.\n"
        "9. Estimar dados ausentes com base em tendências recentes."
    ),
    expected_output="JSON estruturado com dados brutos para todos os elementos solicitados.",
    agent=pesquisador
)

analise_tendencias = Task(
    description=(
        "Analisar dados de {setor} para produzir:\n"
        "1. Análise SWOT com gráfico aranha (fornecer dados numéricos para o gráfico).\n"
        "2. 5 Forças de Porter com gráfico aranha (fornecer dados numéricos para o gráfico).\n"
        "3. Modelo TCO e fórmula LPP.\n"
        "4. Avaliação de riscos e estratégias de mitigação.\n"
        "5. Alavancas de negociação e BATNA.\n"
        "6. Momento ideal para negociação com base em sazonalidade."
    ),
    # IMPORTANT: The expected output should ideally be structured (e.g., JSON)
    # containing the analysis results AND the data needed for the graphs.
    # The current implementation uses static data for graphs.
    expected_output="Texto em Markdown com insights analíticos. Os dados para os gráficos SWOT e Porter serão gerados separadamente com valores de exemplo.",
    agent=analista
)

redacao_relatorio = Task(
    description=(
        f"Gerar um relatório de {CONFIG_RELATORIO['contagem_palavras']} palavras em Markdown para {{setor}}, "
        f"cobrindo todas as seções: {', '.join(CONFIG_RELATORIO['secoes'])}.\n"
        "Incluir tabelas e recomendações acionáveis.\n"
        # "Exportar para Word com identidade visual da Vorätte.\n" # Export handled outside the task
        f"Garantir que o relatório seja em português, detalhado, com no mínimo {CONFIG_RELATORIO['paginas_minimas']} páginas.\n"
        "Basear-se em dados atuais, com projeções para 12-18 meses.\n"
        "O output deve ser apenas o conteúdo Markdown completo do relatório."
    ),
    expected_output="String contendo o relatório completo em formato Markdown.",
    agent=redator
)

# Equipe
equipe = Crew(
    agents=[pesquisador, analista, redator],
    tasks=[coleta_dados, analise_tendencias, redacao_relatorio],
    verbose=2 # Use verbose=2 for more detailed logs
)

# Função principal
def gerar_relatorio(setor, regiao, empresa):
    # Define unique filenames for this run to avoid conflicts if run concurrently
    run_id = uuid.uuid4()
    swot_image_file = f'swot_aranha_{run_id}.png'
    porter_image_file = f'porter_aranha_{run_id}.png'
    arquivo_saida_base = f"relatorio_{setor.replace(' ', '_').replace('/', '-')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    arquivo_saida_docx = f"{arquivo_saida_base}.docx"
    # arquivo_saida_md = f"{arquivo_saida_base}.md" # Optional: save markdown too

    try:
        logging.info(f"Iniciando kickoff da equipe para Setor: {setor}, Região: {regiao}, Empresa: {empresa}")
        resultado_crew = equipe.kickoff(inputs={
            "setor": setor,
            "regiao": regiao,
            "empresa": empresa
        })
        logging.info("Kickoff da equipe concluído.")

        # Extrair o resultado final (o output da última task, que deve ser o Markdown)
        # Acessar resultado.raw pode ser frágil; verificar a estrutura de `resultado_crew`
        if hasattr(resultado_crew, 'raw'):
             conteudo_markdown = str(resultado_crew.raw)
        elif isinstance(resultado_crew, str):
             conteudo_markdown = resultado_crew
        else:
             logging.warning("Não foi possível extrair o conteúdo Markdown do resultado do CrewAI. Usando representação string.")
             conteudo_markdown = str(resultado_crew)

        # Optional: Save the raw markdown output
        # try:
        #     with open(arquivo_saida_md, 'w', encoding='utf-8') as f:
        #         f.write(conteudo_markdown)
        #     logging.info(f"Relatório Markdown salvo como {arquivo_saida_md}")
        # except Exception as md_e:
        #     logging.error(f"Erro ao salvar relatório Markdown: {md_e}")

        # --- Geração de Gráficos (com dados de exemplo, conforme solicitado) ---
        logging.info("Gerando gráficos de exemplo...")
        swot_dados_exemplo = {"Escalabilidade": 9, "Custo": 6, "Inovação": 8, "Risco": 5}
        porter_dados_exemplo = {
            "Poder Fornecedores": 5,
            "Poder Compradores": 8,
            "Novos Entrantes": 3,
            "Substitutos": 6,
            "Rivalidade": 9
        }
        swot_gerado = criar_grafico_aranha(swot_dados_exemplo, swot_dados_exemplo.keys(), "Análise SWOT (Exemplo)", swot_image_file)
        porter_gerado = criar_grafico_aranha(porter_dados_exemplo, porter_dados_exemplo.keys(), "5 Forças de Porter (Exemplo)", porter_image_file)
        # ----------------------------------------------------------------------

        logging.info(f"Exportando relatório para Word: {arquivo_saida_docx}")
        exportar_para_word(conteudo_markdown, setor, arquivo_saida_docx)

        # Limpeza dos arquivos de imagem temporários
        if swot_gerado and os.path.exists(swot_image_file):
            try:
                os.remove(swot_image_file)
                logging.info(f"Arquivo de imagem temporário removido: {swot_image_file}")
            except OSError as rm_e:
                logging.warning(f"Não foi possível remover o arquivo de imagem temporário {swot_image_file}: {rm_e}")
        if porter_gerado and os.path.exists(porter_image_file):
             try:
                os.remove(porter_image_file)
                logging.info(f"Arquivo de imagem temporário removido: {porter_image_file}")
             except OSError as rm_e:
                logging.warning(f"Não foi possível remover o arquivo de imagem temporário {porter_image_file}: {rm_e}")

        return arquivo_saida_docx

    except Exception as e:
        logging.error(f"Erro geral ao gerar relatório para {setor}: {e}", exc_info=True)
        # Limpar arquivos temporários em caso de erro também
        if os.path.exists(swot_image_file):
            try: os.remove(swot_image_file) 
            except OSError: pass
        if os.path.exists(porter_image_file):
            try: os.remove(porter_image_file)
            except OSError: pass
        raise # Re-lança a exceção para ser tratada pelo Streamlit

# Bloco de execução principal (para testes locais, não usado pelo Streamlit)
if __name__ == "__main__":
    logging.info("Executando script analisemercado.py diretamente para teste...")
    # Certifique-se de que as variáveis de ambiente estão definidas para teste local
    # (OPENAI_API_KEY, SERPER_API_KEY, GROQ_API_KEY)
    if not os.getenv("OPENAI_API_KEY") or not os.getenv("SERPER_API_KEY") or not os.getenv("GROQ_API_KEY"):
        print("ERRO: Variáveis de ambiente (OPENAI_API_KEY, SERPER_API_KEY, GROQ_API_KEY) não definidas para teste local.")
    else:
        try:
            arquivo_gerado = gerar_relatorio("Serviços de Computação em Nuvem", "Brasil", "Grande Empresa de Teste")
            print(f"Teste concluído. Relatório gerado: {arquivo_gerado}")
        except Exception as main_e:
            print(f"Erro durante a execução do teste: {main_e}")

